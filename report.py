# ============================================================
# Word (.docx) export.
#
# Produces a clinician-ready document laid out like a neuropsychological
# report annex:
#   1. Title block (identifier, date).
#   2. Visual profile: every radar in one compact grid, at most two
#      rows, at the top of the page.
#   3. Band legend and the percentile table (one continuous table with
#      shaded domain sections; one row per data series per measure).
#   4. Clinical notes (per domain and global), when provided.
#   5. Interpretive draft.
#   6. Optional lexicon of the assessed functions.
# Every page carries the clinician's name bottom-left as a discreet
# watermark, with the page number bottom-right.
#
# Arial throughout. Everything is built in memory and written only to
# the path the user picks. No patient data is persisted anywhere else.
# ============================================================

from __future__ import annotations

import io
from datetime import date
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from engine import ProfileResult, classify_band, format_percentile
from plots import (
    band_index,
    close_fig,
    composite_figure,
    fig_to_png_bytes,
    resolve_theme,
)

# --- 0. Palette and bilingual strings -------------------------

FONT = "Arial"
ACCENT_HEX = "2C7290"
ACCENT_DEEP = RGBColor(0x22, 0x5D, 0x77)
INK = RGBColor(0x1D, 0x2B, 0x33)
MUTED = RGBColor(0x6B, 0x77, 0x7E)
STRENGTH = RGBColor(0x1F, 0x7A, 0x63)
WEAKNESS = RGBColor(0x56, 0x58, 0xA6)
# Neutral grays for the table chrome (the band column is the only color).
HEADER_FILL = "E9ECEF"
DOMAIN_FILL = "F2F4F6"
MEAN_FILL = "F8F9FA"
RULE_FILL = "CED5DA"

_METRIC_LABEL = {
    "z": "z", "percentile": "%ile", "scaled": "scaled",
    "standard": "standard", "t": "T",
}

_MONTHS = {
    "fr": ["janvier", "février", "mars", "avril", "mai", "juin", "juillet",
           "août", "septembre", "octobre", "novembre", "décembre"],
    "en": ["January", "February", "March", "April", "May", "June", "July",
           "August", "September", "October", "November", "December"],
}

_STRINGS = {
    "fr": {
        "title": "PROFIL COGNITIF",
        "identifier": "Identifiant", "date": "Date",
        "figures_title": "Profil visuel",
        "legend_title": "Bandes de classification",
        "table_title": "Tableau des percentiles",
        "notes_title": "Notes cliniques",
        "global_note": "Note generale",
        "draft_title": "Interprétation (brouillon)",
        "lexicon_title": "Lexique des fonctions évaluées",
        "cols": ["Sous-fonction", "Score", "Percentile", "Bande", "Indicateur"],
        "domain_mean": "Moyenne du domaine",
        "strength": "▲ Force", "weakness": "▼ Faiblesse", "within": "-",
        "no_data": "Aucune mesure administrée.",
        "footer": ("Document généré localement. Brouillon à réviser et à "
                   "valider selon le jugement clinique."),
    },
    "en": {
        "title": "COGNITIVE PROFILE",
        "identifier": "Identifier", "date": "Date",
        "figures_title": "Visual profile",
        "legend_title": "Classification bands",
        "table_title": "Percentile table",
        "notes_title": "Clinical notes",
        "global_note": "General note",
        "draft_title": "Interpretation (draft)",
        "lexicon_title": "Lexicon of assessed functions",
        "cols": ["Sub-function", "Score", "Percentile", "Band", "Indicator"],
        "domain_mean": "Domain mean",
        "strength": "▲ Strength", "weakness": "▼ Weakness", "within": "-",
        "no_data": "No measure administered.",
        "footer": ("Generated locally. Draft to be reviewed and validated "
                   "with clinical judgement."),
    },
}

# Column widths (sum ~6.9 inches inside 0.8 inch side margins on Letter).
_COLW = [Inches(2.6), Inches(1.15), Inches(0.95), Inches(1.5), Inches(0.7)]


def _today_str(lang: str) -> str:
    """Readable date without depending on the machine locale."""
    today = date.today()
    months = _MONTHS["fr"] if lang == "fr" else _MONTHS["en"]
    if lang == "fr":
        return f"{today.day} {months[today.month - 1]} {today.year}"
    return f"{months[today.month - 1]} {today.day}, {today.year}"


# --- 1. Low-level docx helpers --------------------------------

def _shade(cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _run(paragraph, text, *, bold=False, size=10.0, color=None, italic=False):
    """Add a styled Arial run."""
    r = paragraph.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def _cell(cell, text, *, bold=False, size=9.0, color=INK, align="left",
          fill=None, italic=False) -> None:
    """Replace a cell's content with one styled run, tightly spaced."""
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    _run(para, text, bold=bold, size=size, color=color, italic=italic)
    if fill:
        _shade(cell, fill)


def _compact_cell_margins(table, top=22, bottom=22, left=70, right=70) -> None:
    """Tighten table cell margins (twips) so rows stay compact."""
    tbl_pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tbl_pr.append(mar)


def _set_borders(table, color=RULE_FILL, size=4) -> None:
    """Give the table thin, light borders."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def _rule(paragraph, color=ACCENT_HEX, size=14) -> None:
    """Add a colored bottom border (a horizontal rule) to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def _page_field(paragraph) -> None:
    """Insert a live PAGE number field."""
    run = paragraph.add_run()
    run.font.name = FONT
    run.font.size = Pt(7.5)
    run.font.color.rgb = MUTED
    for kind, txt in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = txt
            run._r.append(instr)


def _heading(doc, text, *, size=12.0, rule=True, space_before=12) -> None:
    """A styled section heading in the accent color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, bold=True, size=size, color=ACCENT_DEEP)
    if rule:
        _rule(p, color=ACCENT_HEX, size=8)


# --- 2. Page sections -----------------------------------------

def _title_block(doc, s, patient_id, lang) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    _run(title, s["title"], bold=True, size=20, color=ACCENT_DEEP)
    _rule(title, color=ACCENT_HEX, size=18)

    info = doc.add_paragraph()
    info.paragraph_format.space_before = Pt(4)
    info.paragraph_format.space_after = Pt(4)
    _run(info, f"{s['identifier']} ", bold=True, size=10, color=MUTED)
    _run(info, f"{patient_id or '-'}    ", size=10, color=INK)
    _run(info, f"{s['date']} ", bold=True, size=10, color=MUTED)
    _run(info, _today_str(lang), size=10, color=INK)


def _figures(doc, s, profiles, series_labels, lang, options) -> None:
    """Compact radar grid (at most two rows) near the top of the report."""
    fig = composite_figure(profiles, series_labels, lang,
                           options.get("radial_mode", "z"),
                           options.get("show_summary", True),
                           options.get("theme"))
    fig_w, fig_h = fig.get_size_inches()
    png = fig_to_png_bytes(fig, dpi=300)
    close_fig(fig)
    # Keep the grid clearly in the upper part of the page.
    width_in = min(6.9, 5.7 * fig_w / fig_h)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.add_run().add_picture(io.BytesIO(png), width=Inches(width_in))


def _legend(doc, s, lang, bands) -> None:
    labels_fr = ["Extr. bas", "Limite", "Moy. inf.", "Moyenne",
                 "Moy. sup.", "Supérieur", "Très sup."]
    labels_en = ["Extr. low", "Borderline", "Low avg", "Average",
                 "High avg", "Superior", "Very sup."]
    labels = labels_fr if lang == "fr" else labels_en
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _compact_cell_margins(table, top=14, bottom=14, left=40, right=40)
    for i, cell in enumerate(table.rows[0].cells):
        _cell(cell, labels[i], size=7, align="center", fill=bands[i])
        cell.width = Inches(0.98)


def _table(doc, s, profiles: list[ProfileResult], series_labels, inputs,
           lang, bands) -> None:
    """One continuous percentile table; with two series, each measure
    carries one row per administered series (series label in the Score
    column), and each domain one mean row per series."""
    _heading(doc, s["table_title"])
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_borders(table)
    _compact_cell_margins(table)

    for i, label in enumerate(s["cols"]):
        align = "left" if i == 0 else "center"
        _cell(table.rows[0].cells[i], label, bold=True, size=8.5, color=INK,
              align=align, fill=HEADER_FILL)
        table.rows[0].cells[i].width = _COLW[i]

    nser = len(profiles)
    multi = nser > 1
    # Per-series cursors into the engine results (administered only).
    cursors = [[0] * len(p.domains) for p in profiles]

    any_data = False
    ndom = len(profiles[0].domains)
    for di in range(ndom):
        domain_rows = []
        nmeas = len(inputs[di]) if di < len(inputs) else 0
        for mi in range(nmeas):
            for si in range(nser):
                raw = inputs[di][mi][si] if si < len(inputs[di][mi]) else None
                if raw is None:
                    continue
                r = profiles[si].domains[di].measures[cursors[si][di]]
                cursors[si][di] += 1
                domain_rows.append((mi, si, raw, r))
        if not domain_rows:
            continue
        any_data = True
        dom = profiles[0].domains[di]

        drow = table.add_row()
        merged = drow.cells[0].merge(drow.cells[4])
        _cell(merged, dom.name(lang), bold=True, size=9, color=INK,
              fill=DOMAIN_FILL)

        last_mi = None
        for mi, si, raw, r in domain_rows:
            cells = table.add_row().cells
            name = r.name(lang) if mi != last_mi else ""
            last_mi = mi
            metric = _METRIC_LABEL.get(raw["metric"], raw["metric"])
            score = f"{raw['value']} ({metric})"
            if multi:
                score = f"{series_labels[si]} · {score}"
            _cell(cells[0], name, size=9)
            _cell(cells[1], score, size=8.5, align="center")
            _cell(cells[2], r.percentile_display, size=9, align="center",
                  bold=True)
            _cell(cells[3], r.band(lang), size=8.5,
                  fill=bands[band_index(r.percentile)])
            marker = {"strength": s["strength"], "weakness": s["weakness"],
                      "within": s["within"]}[r.flag]
            mcolor = {"strength": STRENGTH, "weakness": WEAKNESS,
                      "within": MUTED}[r.flag]
            _cell(cells[4], marker, size=8, align="center", color=mcolor)
            for i in range(5):
                cells[i].width = _COLW[i]

        for si in range(nser):
            d = profiles[si].domains[di]
            if d.mean_percentile is None:
                continue
            cells = table.add_row().cells
            label = s["domain_mean"]
            if multi:
                label = f"{label} ({series_labels[si]})"
            _cell(cells[0], label, bold=True, size=8.5, fill=MEAN_FILL)
            _cell(cells[1], "", size=8.5, fill=MEAN_FILL)
            _cell(cells[2], format_percentile(d.mean_percentile), bold=True,
                  size=8.5, align="center", fill=MEAN_FILL)
            band_fr, band_en = classify_band(d.mean_percentile)
            _cell(cells[3], band_fr if lang == "fr" else band_en, bold=True,
                  size=8.5, fill=bands[band_index(d.mean_percentile)])
            _cell(cells[4], "", size=8.5, fill=MEAN_FILL)
            for i in range(5):
                cells[i].width = _COLW[i]

    if not any_data:
        doc.add_paragraph(s["no_data"])


def _notes(doc, s, profiles, lang, options) -> None:
    """Clinical notes: one entry per annotated domain, plus a global
    note. Section omitted entirely when no note was written."""
    notes = options.get("notes") or {}
    per_domain = notes.get("domains") or []
    global_note = str(notes.get("global") or "").strip()
    entries = []
    for di, d in enumerate(profiles[0].domains):
        raw = per_domain[di] if di < len(per_domain) else ""
        text = str(raw or "").strip()
        if text:
            entries.append((d.name(lang), text))
    if not entries and not global_note:
        return
    _heading(doc, s["notes_title"])
    for name, text in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, f"{name}. ", bold=True, size=10, color=INK)
        _run(p, text, size=10, color=INK)
    if global_note:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, f"{s['global_note']}. ", bold=True, size=10, color=INK)
        _run(p, global_note, size=10, color=INK)


def _interpretation(doc, s, draft_text, lang) -> None:
    _heading(doc, s["draft_title"])
    first = True
    for line in (draft_text or "").split("\n"):
        if line.strip() == "":
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(5)
        if first:
            _run(p, line, size=9, italic=True, color=MUTED)
            first = False
        else:
            _run(p, line, size=10.5, color=INK)


def _lexicon(doc, s, options) -> None:
    """Optional lexicon: term + definition lines, in the export language
    (the UI assembles the checked, possibly edited entries)."""
    items = options.get("lexicon") or []
    items = [it for it in items
             if str(it.get("term", "")).strip()
             and str(it.get("definition", "")).strip()]
    if not items:
        return
    _heading(doc, s["lexicon_title"])
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _run(p, f"{it['term'].strip()}. ", bold=True, size=9.5, color=INK)
        _run(p, it["definition"].strip(), size=9.5, color=INK)


def _footer(doc, s, clinician: str) -> None:
    """Watermark footer: clinician bottom-left, page number bottom-right,
    and the local-generation disclaimer on a second, smaller line."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = ""
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.9),
                                              WD_TAB_ALIGNMENT.RIGHT)
    name = (clinician or "").strip()
    if name:
        _run(p, name, size=8.5, color=MUTED, bold=True)
    _run(p, "\t", size=8.5)
    _page_field(p)
    p2 = footer.add_paragraph()
    _run(p2, s["footer"], italic=True, size=6.5, color=MUTED)


# --- 3. Document assembly -------------------------------------

def build_document(profiles: list[ProfileResult],
                   series_labels: list[str],
                   inputs: list,
                   draft_text: str,
                   patient_id: str,
                   lang: str = "fr",
                   options: Optional[dict] = None) -> Document:
    """Build and return the full report as a python-docx Document."""
    options = options or {}
    bands, _accent, _accent_deep = resolve_theme(options.get("theme"))
    s = _STRINGS["fr"] if lang == "fr" else _STRINGS["en"]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    any_data = any(d.measures for p in profiles for d in p.domains)

    _title_block(doc, s, patient_id, lang)
    if any_data:
        _figures(doc, s, profiles, series_labels, lang, options)
        _legend(doc, s, lang, bands)
    _table(doc, s, profiles, series_labels, inputs, lang, bands)
    _notes(doc, s, profiles, lang, options)
    _interpretation(doc, s, draft_text, lang)
    _lexicon(doc, s, options)
    _footer(doc, s, options.get("clinician", ""))

    return doc


def export_report(path: str,
                  profiles: list[ProfileResult],
                  series_labels: list[str],
                  inputs: list,
                  draft_text: str,
                  patient_id: str,
                  lang: str = "fr",
                  options: Optional[dict] = None) -> str:
    """Build the document and save it to path. Returns the path."""
    doc = build_document(profiles, series_labels, inputs, draft_text,
                         patient_id, lang, options)
    doc.save(path)
    return path


# --- 4. Manual smoke test ------------------------------------

if __name__ == "__main__":
    import tempfile
    from engine import (DomainInput, MeasureInput, generate_report_text,
                        process_profile)

    demo = [
        DomainInput("Attention / Vitesse", "Attention / Speed", [
            MeasureInput("Soutenue", "Sustained", 8, "scaled"),
            MeasureInput("Sélective", "Selective", 50, "t"),
            MeasureInput("Divisée", "Divided", 5, "percentile"),
            MeasureInput("Vigilance", "Vigilance", -0.5, "z"),
        ]),
        DomainInput("Mémoire", "Memory", [
            MeasureInput("À court terme", "Short-term", 25, "percentile"),
            MeasureInput("À long terme", "Long-term", 1.5, "z"),
            MeasureInput("MdeT auditive", "Auditory WM", 60, "t"),
        ]),
    ]
    demo2 = [
        DomainInput("Attention / Vitesse", "Attention / Speed", [
            MeasureInput("Soutenue", "Sustained", 11, "scaled"),
            MeasureInput("Sélective", "Selective", 55, "t"),
            MeasureInput("Vigilance", "Vigilance", 0.2, "z"),
        ]),
        DomainInput("Mémoire", "Memory", [
            MeasureInput("À court terme", "Short-term", 45, "percentile"),
            MeasureInput("À long terme", "Long-term", 1.1, "z"),
            MeasureInput("MdeT auditive", "Auditory WM", 63, "t"),
        ]),
    ]
    p1 = process_profile(demo, "DEMO-01", 1.0)
    p2 = process_profile(demo2, "DEMO-01", 1.0)
    inputs = [
        [[{"value": "8", "metric": "scaled"}, {"value": "11", "metric": "scaled"}],
         [{"value": "50", "metric": "t"}, {"value": "55", "metric": "t"}],
         [{"value": "5", "metric": "percentile"}, None],
         [{"value": "-0.5", "metric": "z"}, {"value": "0.2", "metric": "z"}]],
        [[{"value": "25", "metric": "percentile"}, {"value": "45", "metric": "percentile"}],
         [{"value": "1.5", "metric": "z"}, {"value": "1.1", "metric": "z"}],
         [{"value": "60", "metric": "t"}, {"value": "63", "metric": "t"}]],
    ]
    draft = generate_report_text(p1, "fr")
    options = {
        "radial_mode": "z",
        "clinician": "Nicola Thibault, PhD.",
        "notes": {"domains": ["Rendement attentionnel variable selon la "
                              "condition de medication.", ""],
                  "global": "Collaboration adequate aux deux temps."},
        "lexicon": [{"term": "Soutenue",
                     "definition": "Capacite a maintenir l'attention sur "
                                   "une tache pendant une periode prolongee."}],
    }
    out = export_report(tempfile.mktemp(suffix=".docx"), [p1, p2],
                        ["Sans médication", "Avec médication"], inputs,
                        draft, "DEMO-01", "fr", options)
    print("Wrote", out)
